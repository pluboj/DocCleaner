from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


class TextExtractor:

    def __init__(self, filename):
        self.filename = filename

    def process_text(self):
        document = Document(self.filename)
        path = self.filename.split('/')
        path[-1] = "DOC.docx"
        file_path = '/'.join(path)
        new_doc = Document()

        for block in self.iter_block_items(document):

            # print(block.text if isinstance(block, Paragraph) else '<table>')
            if isinstance(block, Paragraph):
                formatted_block = self.add_markups(block.runs)
                if formatted_block is not "":
                    new_doc.add_paragraph(formatted_block)
            elif isinstance(block, Table):
                for row in block.rows:
                    row_data = []
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            formatted_para = self.add_markups(paragraph.runs)
                            if formatted_para is not "":
                                new_doc.add_paragraph(formatted_para)

        new_doc.save(file_path)

    @staticmethod
    def iter_block_items(parent):

        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        elif isinstance(parent, _Row):
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def add_markups(self, chunk):
        item = []
        for run in chunk:
            if len(run.text) == 1 and run.text.upper() == 'O':
                continue
            else:
                formatted_run = run.text
                if run.underline:
                    formatted_run = "<u>" + formatted_run + "</u>"
                if run.bold:
                    formatted_run = "<strong>" + formatted_run + "</strong>"
                if run.italic:
                    formatted_run = "<em>" + formatted_run + "</em>"

            item.append(formatted_run)
        joined_item = "".join(item)
        joined_item = joined_item.replace("</strong><strong>", "").replace("</em><em>", "") \
            .replace("</u><u>", "")
        if len(joined_item) == 1 and joined_item.upper() == 'O':
            return ''
        return joined_item
